"""
文件解析器 —— 将不同格式的文件转换为统一的纯文本 + 段落列表
"""
import os
import re
from typing import Tuple, List
from .base_detector import Paragraph


class FileParser:
    """多格式文件解析器"""

    SUPPORTED_FORMATS = {".docx", ".txt", ".pdf"}
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

    def parse(self, filepath: str) -> Tuple[str, List[Paragraph], dict]:
        """
        解析文件，返回 (纯文本, 段落列表, 文件元信息)
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"文件不存在: {filepath}")

        file_size = os.path.getsize(filepath)
        if file_size > self.MAX_FILE_SIZE:
            raise ValueError(f"文件大小 {file_size / 1024 / 1024:.1f}MB 超过限制 5MB")
        if file_size == 0:
            raise ValueError("文件为空")

        ext = os.path.splitext(filepath)[1].lower()
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的文件格式: {ext}，支持格式: {', '.join(self.SUPPORTED_FORMATS)}")

        filename = os.path.basename(filepath)
        meta = {
            "filename": filename,
            "format": ext.replace(".", ""),
            "size": file_size,
            "size_display": self._format_size(file_size),
        }

        if ext == ".txt":
            full_text, paragraphs = self._parse_txt(filepath)
        elif ext == ".docx":
            full_text, paragraphs = self._parse_docx(filepath)
        elif ext == ".pdf":
            full_text, paragraphs = self._parse_pdf(filepath)
        else:
            raise ValueError(f"不支持的格式: {ext}")

        # 清理文本
        full_text = self._clean_text(full_text)
        meta["char_count"] = len(full_text)
        meta["paragraph_count"] = len(paragraphs)

        return full_text, paragraphs, meta

    def _parse_txt(self, filepath: str) -> Tuple[str, List[Paragraph]]:
        """解析 TXT 文件"""
        content = None
        for encoding in ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]:
            try:
                with open(filepath, "r", encoding=encoding) as f:
                    content = f.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if content is None:
            raise ValueError("无法识别文件编码，请确保文件为 UTF-8 或 GBK 编码")

        return self._text_to_paragraphs(content)

    def _parse_docx(self, filepath: str) -> Tuple[str, List[Paragraph]]:
        """解析 DOCX 文件"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        try:
            doc = Document(filepath)
        except Exception as e:
            raise ValueError(f"DOCX 文件解析失败: {str(e)}")

        texts = []
        paragraphs = []
        char_offset = 0

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append(Paragraph(
                    index=len(paragraphs) + 1,
                    page=None,
                    text=text,
                    start_char=char_offset,
                ))
                texts.append(text)
                char_offset += len(text) + 1  # +1 for newline

        full_text = "\n".join(texts)
        return full_text, paragraphs

    def _parse_pdf(self, filepath: str) -> Tuple[str, List[Paragraph]]:
        """解析 PDF 文件"""
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                raise ImportError("请安装 pypdf: pip install pypdf")

        try:
            reader = PdfReader(filepath)
        except Exception as e:
            raise ValueError(f"PDF 文件解析失败: {str(e)}")

        texts = []
        paragraphs = []
        char_offset = 0

        for page_num, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""

            # 按段落拆分
            page_paragraphs = [p.strip() for p in page_text.split("\n") if p.strip()]

            for para_text in page_paragraphs:
                paragraphs.append(Paragraph(
                    index=len(paragraphs) + 1,
                    page=page_num,
                    text=para_text,
                    start_char=char_offset,
                ))
                texts.append(para_text)
                char_offset += len(para_text) + 1

        full_text = "\n".join(texts)
        if not full_text.strip():
            raise ValueError("PDF 文件无法提取文本内容（可能是扫描件，当前版本暂不支持 OCR）")

        return full_text, paragraphs

    def _text_to_paragraphs(self, content: str) -> Tuple[str, List[Paragraph]]:
        """将纯文本按段落拆分"""
        lines = content.split("\n")
        paragraphs = []
        texts = []
        char_offset = 0

        for line in lines:
            text = line.strip()
            if text:
                paragraphs.append(Paragraph(
                    index=len(paragraphs) + 1,
                    page=None,
                    text=text,
                    start_char=char_offset,
                ))
                texts.append(text)
            char_offset += len(line) + 1

        full_text = "\n".join(texts)
        return full_text, paragraphs

    def _clean_text(self, text: str) -> str:
        """清理文本：统一全角/半角，去多余空白"""
        # 统一全角数字为半角
        result = []
        for ch in text:
            code = ord(ch)
            if 0xFF10 <= code <= 0xFF19:  # ０-９
                result.append(chr(code - 0xFF10 + 0x30))
            elif 0xFF21 <= code <= 0xFF3A:  # Ａ-Ｚ
                result.append(chr(code - 0xFF21 + 0x41))
            elif 0xFF41 <= code <= 0xFF5A:  # ａ-ｚ
                result.append(chr(code - 0xFF41 + 0x61))
            else:
                result.append(ch)
        text = "".join(result)
        # 去除多余空行
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text

    def _format_size(self, size_bytes: int) -> str:
        """格式化文件大小"""
        if size_bytes < 1024:
            return f"{size_bytes}B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f}KB"
        else:
            return f"{size_bytes / 1024 / 1024:.1f}MB"

    def extract_and_ocr(self, filepath: str) -> List[dict]:
        """
        提取图片并使用 MinerU 进行 OCR 识别
        返回: List[dict] - 每项包含 {page, image_path, ocr_text, text_lines}
        """
        ext = os.path.splitext(filepath)[1].lower()
        results = []

        if ext == ".pdf":
            results = self._ocr_pdf_images(filepath)
        elif ext == ".docx":
            results = self._ocr_docx_images(filepath)
        elif ext == ".txt":
            pass

        return results

    def _ocr_pdf_images(self, filepath: str) -> List[dict]:
        """使用 MinerU 对 PDF 中的图片进行 OCR 识别"""
        results = []

        try:
            from magic_pdf.data.data_reader_writer import FileBasedDataReader
            from magic_pdf.data.dataset import PymuDocDataset
            from magic_pdf.model import UNIPipe
            import json
            import tempfile
            import shutil
        except ImportError:
            return self._ocr_pdf_fallback(filepath)

        try:
            temp_dir = tempfile.mkdtemp()
            image_dir = os.path.join(temp_dir, "images")
            os.makedirs(image_dir, exist_ok=True)

            reader = FileBasedDataReader("")

            with open(filepath, "rb") as f:
                pdf_bytes = f.read()

            ds = PymuDocDataset(pdf_bytes, None)

            if ds.classify() == "pdf":
                pipe = UNIPipe(ds, {"_pdf_type": "", "model_list": "auto"})
                pipe.pipe_classify()
                pipe.pipe_parse()

                content_list = pipe.get_content_list(FileBasedDataReader(""))

                for item in content_list:
                    page_num = item.get("page_idx", 0) + 1
                    ocr_text = item.get("text", "")

                    if ocr_text and ocr_text.strip():
                        text_lines = [line.strip() for line in ocr_text.split("\n") if line.strip()]
                        results.append({
                            "page": page_num,
                            "image_path": None,
                            "ocr_text": ocr_text,
                            "text_lines": text_lines,
                            "source": "pdf_ocr"
                        })

            shutil.rmtree(temp_dir, ignore_errors=True)

        except Exception as e:
            return self._ocr_pdf_fallback(filepath)

        return results

    def _ocr_pdf_fallback(self, filepath: str) -> List[dict]:
        """使用 pdf2image + PaddleOCR 的后备方案"""
        results = []

        try:
            from pdf2image import convert_from_path
            import paddle
            from paddleocr import PaddleOCR
        except ImportError:
            return results

        try:
            paddle.enable_static()
            ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)

            images = convert_from_path(filepath)

            for page_num, image in enumerate(images, 1):
                result = ocr.ocr(image, cls=True)

                if result and result[0]:
                    ocr_text_lines = []
                    for line in result[0]:
                        if line and len(line) >= 2:
                            text = line[1][0]
                            if text.strip():
                                ocr_text_lines.append(text)

                    if ocr_text_lines:
                        results.append({
                            "page": page_num,
                            "image_path": None,
                            "ocr_text": "\n".join(ocr_text_lines),
                            "text_lines": ocr_text_lines,
                            "source": "pdf_ocr_fallback"
                        })

        except Exception:
            pass

        return results

    def _ocr_docx_images(self, filepath: str) -> List[dict]:
        """对 DOCX 中的图片进行 OCR 识别"""
        results = []

        try:
            from docx import Document
            from docx.parts.inline_shapes import InlineShapes
            import io
            from PIL import Image
            import paddle
            from paddleocr import PaddleOCR
        except ImportError:
            return results

        try:
            paddle.enable_static()
            ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)

            doc = Document(filepath)

            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_part = rel.target_part
                    image = Image.open(io.BytesIO(image_part.blob))

                    result = ocr.ocr(image, cls=True)

                    if result and result[0]:
                        ocr_text_lines = []
                        for line in result[0]:
                            if line and len(line) >= 2:
                                text = line[1][0]
                                if text.strip():
                                    ocr_text_lines.append(text)

                        if ocr_text_lines:
                            results.append({
                                "page": 1,
                                "image_path": None,
                                "ocr_text": "\n".join(ocr_text_lines),
                                "text_lines": ocr_text_lines,
                                "source": "docx_image_ocr"
                            })

        except Exception:
            pass

        return results
